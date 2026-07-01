"""
Export Service for generating PDF and DOCX files from book content
"""
import io
import asyncio
import html
from typing import Dict, List, Optional, BinaryIO
from datetime import datetime
import re

from app.services.export_templates import PAGE_SIZES_INCHES, resolve_template

# Optional export dependencies are guarded so the service (and the rest of the
# API) still imports if a library is missing — only the affected format fails,
# and it fails loudly with a clear message rather than at import time.
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak,
        Table, TableStyle, KeepTogether
    )
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def _add_page_number_field(paragraph) -> None:
    """Append a Word PAGE field to a paragraph (live page number in footer)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)

try:
    from ebooklib import epub
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False

try:
    import html2text
    HTML2TEXT_AVAILABLE = True
except ImportError:
    HTML2TEXT_AVAILABLE = False


class ExportUnavailableError(RuntimeError):
    """Raised when an export format's required library is not installed."""


class ExportValidationError(ValueError):
    """Raised when a book has no exportable content (clear, user-facing message)."""


class ExportTimeoutError(TimeoutError):
    """Raised when export generation exceeds the configured time budget."""


class ExportService:
    """Service for exporting books to PDF and DOCX formats."""

    def __init__(self):
        if HTML2TEXT_AVAILABLE:
            self.h2t = html2text.HTML2Text()
            self.h2t.ignore_links = False
            self.h2t.ignore_images = True
            self.h2t.body_width = 0  # Don't wrap lines
        else:
            self.h2t = None

    def _clean_html_content(self, content: str) -> str:
        """Convert HTML content to clean text, preserving basic formatting."""
        if not content:
            return ""

        if self.h2t is not None:
            # Convert HTML to markdown
            markdown_text = self.h2t.handle(content)
        else:
            # ponytail: html2text not installed — fall back to a tag strip that
            # preserves block boundaries so words/paragraphs don't merge.
            text = re.sub(r'(?is)<(script|style).*?>.*?</\1>', '', content)
            text = re.sub(r'(?i)</(p|div|h[1-6]|li|br|tr)>', '\n', text)
            text = re.sub(r'(?is)<[^>]+>', '', text)
            markdown_text = html.unescape(text)

        # Clean up excessive newlines
        markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text)

        return markdown_text.strip()

    def _extract_text_formatting(self, content: str) -> List[Dict]:
        """Extract text and basic formatting from HTML content."""
        # This is a simple implementation - could be enhanced with proper HTML parsing
        clean_text = self._clean_html_content(content)

        # Split into paragraphs
        paragraphs = clean_text.split('\n\n')

        formatted_content = []
        for para in paragraphs:
            if not para.strip():
                continue

            # Detect headings
            if para.startswith('# '):
                formatted_content.append({
                    'text': para[2:].strip(),
                    'style': 'heading1'
                })
            elif para.startswith('## '):
                formatted_content.append({
                    'text': para[3:].strip(),
                    'style': 'heading2'
                })
            elif para.startswith('### '):
                formatted_content.append({
                    'text': para[4:].strip(),
                    'style': 'heading3'
                })
            else:
                # Regular paragraph
                formatted_content.append({
                    'text': para.strip(),
                    'style': 'normal'
                })

        return formatted_content

    @staticmethod
    def _resolve_tokens(text: str, book_data: Dict) -> str:
        """Fill running-header/footer tokens from book-level data.

        ponytail: running heads are book-level (title/author), not per-chapter —
        {page} is handled at draw time since only the canvas knows the number.
        """
        if not text:
            return ""
        title = book_data.get("title", "Untitled")
        author = book_data.get("author_name") or book_data.get("owner_name") or ""
        return (
            text.replace("{book_title}", title)
            .replace("{running_head}", title)
            .replace("{author}", author)
            .replace("{chapter_title}", title)
        )

    def _make_header_footer(self, template: Dict, book_data: Dict, page):
        """Return an onPage(canvas, doc) callback drawing the running header/footer."""
        page_w, page_h = page
        margins = template["margins"]
        left = self._resolve_tokens(template.get("header", {}).get("left", ""), book_data)
        right = self._resolve_tokens(template.get("header", {}).get("right", ""), book_data)
        center = template.get("footer", {}).get("center", "")

        def _draw(canvas, doc):
            canvas.saveState()
            canvas.setFont("Helvetica", 9)
            canvas.setFillGray(0.4)
            y_header = page_h - margins["top"] * inch + 0.25 * inch
            y_footer = margins["bottom"] * inch - 0.35 * inch
            if left:
                canvas.drawString(margins["inside"] * inch, y_header, left)
            if right:
                canvas.drawRightString(page_w - margins["outside"] * inch, y_header, right)
            if center:
                text = center.replace("{page}", str(canvas.getPageNumber()))
                canvas.drawCentredString(page_w / 2.0, y_footer, text)
            canvas.restoreState()

        return _draw

    async def generate_pdf(
        self,
        book_data: Dict,
        chapters: List[Dict],
        output_stream: Optional[BinaryIO] = None,
        page_size: str = "letter",
        template: Optional[Dict] = None,
    ) -> bytes:
        """
        Generate a PDF from book data and chapters.

        ReportLab is synchronous and CPU-bound, so the build runs in a worker
        thread to keep the event loop responsive for large books and to let an
        outer asyncio timeout actually cancel a runaway export.

        ``template`` (optional) is a resolved export-template dict; when omitted
        the legacy hardcoded styling is used unchanged.
        """
        return await asyncio.to_thread(
            self._build_pdf, book_data, chapters, output_stream, page_size, template
        )

    def _build_pdf(
        self,
        book_data: Dict,
        chapters: List[Dict],
        output_stream: Optional[BinaryIO] = None,
        page_size: str = "letter",
        template: Optional[Dict] = None,
    ) -> bytes:
        if not PDF_AVAILABLE:
            raise ExportUnavailableError(
                "PDF export unavailable: reportlab is not installed"
            )

        # Create output buffer if not provided
        if output_stream is None:
            output_stream = io.BytesIO()

        # Page size + margins: template wins, else legacy 1" margins on letter/A4.
        if template:
            w_in, h_in = PAGE_SIZES_INCHES[template["page_size"]]
            page = (w_in * inch, h_in * inch)
            m = template["margins"]
            margin_left, margin_right = m["inside"] * inch, m["outside"] * inch
            margin_top, margin_bottom = m["top"] * inch, m["bottom"] * inch
        else:
            page = letter if page_size == "letter" else A4
            margin_left = margin_right = margin_top = margin_bottom = 72

        # Create document
        doc = SimpleDocTemplate(
            output_stream,
            pagesize=page,
            rightMargin=margin_right,
            leftMargin=margin_left,
            topMargin=margin_top,
            bottomMargin=margin_bottom,
        )

        # Create styles
        styles = getSampleStyleSheet()

        # Body font/size/leading/indent from the template, else legacy defaults.
        if template:
            body_font = template["font"]["pdf_font"]
            body_size = template["font"]["size"]
            body_leading = body_size * template["line_height"]
            body_indent = template["first_line_indent"] * inch
            heading_font = (
                "Helvetica-Bold" if template["font"]["family"] == "sans" else "Times-Bold"
            )
        else:
            body_font, body_size, body_leading, body_indent = "Helvetica", 11, 16, 0
            heading_font = "Helvetica-Bold"

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName=heading_font,
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Normal'],
            fontSize=16,
            textColor=colors.HexColor('#666666'),
            spaceAfter=12,
            alignment=TA_CENTER
        )

        chapter_title_style = ParagraphStyle(
            'ChapterTitle',
            parent=styles['Heading1'],
            fontName=heading_font,
            fontSize=20,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=20,
            spaceBefore=30
        )

        heading2_style = ParagraphStyle(
            'Heading2',
            parent=styles['Heading2'],
            fontName=heading_font,
            fontSize=16,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=12,
            spaceBefore=20
        )

        heading3_style = ParagraphStyle(
            'Heading3',
            parent=styles['Heading3'],
            fontName=heading_font,
            fontSize=14,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=10,
            spaceBefore=15
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=body_font,
            fontSize=body_size,
            alignment=TA_JUSTIFY,
            spaceAfter=12,
            leading=body_leading,
            firstLineIndent=body_indent,
        )

        # Build story
        story = []

        # Title page
        story.append(Paragraph(book_data.get('title', 'Untitled'), title_style))

        if book_data.get('subtitle'):
            story.append(Paragraph(book_data['subtitle'], subtitle_style))

        story.append(Spacer(1, 0.5*inch))

        # Author info
        if book_data.get('author_name'):
            story.append(Paragraph(f"by {book_data['author_name']}", subtitle_style))

        # Genre and audience
        metadata_parts = []
        if book_data.get('genre'):
            metadata_parts.append(f"Genre: {book_data['genre']}")
        if book_data.get('target_audience'):
            metadata_parts.append(f"Target Audience: {book_data['target_audience']}")

        if metadata_parts:
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(' • '.join(metadata_parts), styles['Normal']))

        # Description
        if book_data.get('description'):
            story.append(Spacer(1, 0.5*inch))
            story.append(Paragraph(book_data['description'], body_style))

        # Page break after title page
        story.append(PageBreak())

        # Table of Contents
        story.append(Paragraph("Table of Contents", chapter_title_style))
        story.append(Spacer(1, 0.2*inch))

        toc_data = []
        for i, chapter in enumerate(chapters, 1):
            chapter_title = chapter.get('title', f'Chapter {i}')
            # Simple TOC without page numbers for now
            toc_data.append([f"{i}.", chapter_title])

        if toc_data:
            toc_table = Table(toc_data, colWidths=[0.5*inch, 5*inch])
            toc_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ]))
            story.append(toc_table)

        story.append(PageBreak())

        # Chapters
        for i, chapter in enumerate(chapters, 1):
            # Chapter title
            chapter_title = chapter.get('title', f'Chapter {i}')
            story.append(Paragraph(f"Chapter {i}: {chapter_title}", chapter_title_style))

            # Chapter description
            if chapter.get('description'):
                story.append(Paragraph(chapter['description'], styles['Italic']))
                story.append(Spacer(1, 0.2*inch))

            # Chapter content
            content = chapter.get('content', '')
            if content:
                # Extract and format content
                formatted_content = self._extract_text_formatting(content)

                for para in formatted_content:
                    if para['style'] == 'heading1':
                        story.append(Paragraph(para['text'], chapter_title_style))
                    elif para['style'] == 'heading2':
                        story.append(Paragraph(para['text'], heading2_style))
                    elif para['style'] == 'heading3':
                        story.append(Paragraph(para['text'], heading3_style))
                    else:
                        story.append(Paragraph(para['text'], body_style))
            else:
                story.append(Paragraph("(No content yet)", styles['Italic']))

            # Add page break after each chapter except the last
            if i < len(chapters):
                story.append(PageBreak())

        # Build PDF. With a template, draw a running header/footer on every page
        # except the title page (clean first page is the professional convention).
        if template:
            draw = self._make_header_footer(template, book_data, page)
            doc.build(story, onLaterPages=draw)
        else:
            doc.build(story)

        # Get bytes if using BytesIO
        if isinstance(output_stream, io.BytesIO):
            output_stream.seek(0)
            return output_stream.getvalue()

        return b''

    async def generate_docx(
        self,
        book_data: Dict,
        chapters: List[Dict],
        output_stream: Optional[BinaryIO] = None,
        template: Optional[Dict] = None,
    ) -> bytes:
        """
        Generate a DOCX file from book data and chapters.

        python-docx is synchronous and CPU-bound; the build runs in a worker
        thread so large books don't block the event loop and timeouts can fire.

        ``template`` (optional) is a resolved export-template dict; when omitted
        the legacy Word styling is used unchanged.
        """
        return await asyncio.to_thread(
            self._build_docx, book_data, chapters, output_stream, template
        )

    def _apply_docx_template(self, doc, template: Dict, book_data: Dict) -> None:
        """Apply page size, margins, body font/spacing, and running header/footer."""
        section = doc.sections[0]
        w_in, h_in = PAGE_SIZES_INCHES[template["page_size"]]
        section.page_width = Inches(w_in)
        section.page_height = Inches(h_in)
        m = template["margins"]
        section.top_margin = Inches(m["top"])
        section.bottom_margin = Inches(m["bottom"])
        section.left_margin = Inches(m["inside"])
        section.right_margin = Inches(m["outside"])

        # Body font + line spacing on the Normal style, and the template font on
        # the heading styles so titles/TOC/chapter headings inherit it too.
        font_name = template["font"]["docx_font"]
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(template["font"]["size"])
        normal.paragraph_format.line_spacing = template["line_height"]
        for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
            try:
                doc.styles[style_name].font.name = font_name
            except KeyError:
                pass  # style not present in this document's base template

        # Running header (book/author) — page number lives in the footer. A right
        # tab stop at the content width pins the right-hand value to the margin.
        left = self._resolve_tokens(template.get("header", {}).get("left", ""), book_data)
        right = self._resolve_tokens(template.get("header", {}).get("right", ""), book_data)
        if left or right:
            hp = section.header.paragraphs[0]
            if right:
                content_width = Inches(w_in - m["inside"] - m["outside"])
                hp.paragraph_format.tab_stops.add_tab_stop(
                    content_width, WD_TAB_ALIGNMENT.RIGHT
                )
                hp.text = f"{left}\t{right}"
            else:
                hp.text = left
        if template.get("footer", {}).get("center"):
            fp = section.footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_page_number_field(fp)

    def _build_docx(
        self,
        book_data: Dict,
        chapters: List[Dict],
        output_stream: Optional[BinaryIO] = None,
        template: Optional[Dict] = None,
    ) -> bytes:
        if not DOCX_AVAILABLE:
            raise ExportUnavailableError(
                "DOCX export unavailable: python-docx is not installed"
            )

        # Create document
        doc = Document()

        if template:
            self._apply_docx_template(doc, template, book_data)

        # Set up styles
        styles = doc.styles

        # Title page
        title = doc.add_heading(book_data.get('title', 'Untitled'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if book_data.get('subtitle'):
            subtitle = doc.add_paragraph(book_data['subtitle'])
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(18)
            subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)

        doc.add_paragraph()  # Empty line

        # Author
        if book_data.get('author_name'):
            author = doc.add_paragraph(f"by {book_data['author_name']}")
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author.runs[0].font.size = Pt(14)

        # Metadata
        metadata_parts = []
        if book_data.get('genre'):
            metadata_parts.append(f"Genre: {book_data['genre']}")
        if book_data.get('target_audience'):
            metadata_parts.append(f"Target Audience: {book_data['target_audience']}")

        if metadata_parts:
            doc.add_paragraph()
            metadata = doc.add_paragraph(' • '.join(metadata_parts))
            metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Description
        if book_data.get('description'):
            doc.add_paragraph()
            doc.add_paragraph(book_data['description'])

        # Page break
        doc.add_page_break()

        # Table of Contents
        doc.add_heading('Table of Contents', 1)

        for i, chapter in enumerate(chapters, 1):
            chapter_title = chapter.get('title', f'Chapter {i}')
            toc_entry = doc.add_paragraph(f"{i}. {chapter_title}")
            toc_entry.paragraph_format.left_indent = Inches(0.25)

        doc.add_page_break()

        # Chapters
        for i, chapter in enumerate(chapters, 1):
            # Chapter heading
            chapter_title = chapter.get('title', f'Chapter {i}')
            doc.add_heading(f"Chapter {i}: {chapter_title}", 1)

            # Chapter description
            if chapter.get('description'):
                desc = doc.add_paragraph(chapter['description'])
                desc.runs[0].italic = True

            # Chapter content
            content = chapter.get('content', '')
            if content:
                # Extract and format content
                formatted_content = self._extract_text_formatting(content)

                for para in formatted_content:
                    if para['style'] == 'heading1':
                        doc.add_heading(para['text'], 1)
                    elif para['style'] == 'heading2':
                        doc.add_heading(para['text'], 2)
                    elif para['style'] == 'heading3':
                        doc.add_heading(para['text'], 3)
                    else:
                        p = doc.add_paragraph(para['text'])
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                doc.add_paragraph("(No content yet)")
                doc.paragraphs[-1].runs[0].italic = True

            # Add page break after each chapter except the last
            if i < len(chapters):
                doc.add_page_break()

        # Save to stream
        if output_stream is None:
            output_stream = io.BytesIO()

        doc.save(output_stream)

        # Get bytes if using BytesIO
        if isinstance(output_stream, io.BytesIO):
            output_stream.seek(0)
            return output_stream.getvalue()

        return b''

    async def generate_epub(
        self,
        book_data: Dict,
        chapters: List[Dict],
    ) -> bytes:
        """
        Generate an EPUB (3.0) file from book data and chapters.

        ebooklib is synchronous, so the build runs in a worker thread to keep the
        event loop responsive and let an outer asyncio timeout cancel a runaway
        export — same pattern as PDF/DOCX.
        """
        return await asyncio.to_thread(self._build_epub, book_data, chapters)

    def _chapter_to_xhtml(self, chapter: Dict, index: int) -> str:
        """Render a chapter to a complete, well-formed XHTML document.

        Reuses the shared HTML→formatted-text pipeline (same as PDF/DOCX) so the
        output is guaranteed valid XHTML — raw TipTap HTML isn't always XML-valid
        and would break ereaders.
        """
        title = chapter.get("title", f"Chapter {index}")
        parts = [f"<h1>Chapter {index}: {html.escape(title)}</h1>"]

        if chapter.get("description"):
            parts.append(f"<p><em>{html.escape(chapter['description'])}</em></p>")

        content = chapter.get("content", "")
        if content:
            heading_tag = {"heading1": "h2", "heading2": "h3", "heading3": "h4"}
            for para in self._extract_text_formatting(content):
                text = html.escape(para["text"])
                tag = heading_tag.get(para["style"], "p")
                parts.append(f"<{tag}>{text}</{tag}>")
        else:
            parts.append("<p><em>(No content yet)</em></p>")

        # No <?xml?> prolog: ebooklib parses content as a str during nav
        # generation, and lxml rejects a unicode string carrying an encoding
        # declaration. ebooklib writes its own prolog to the EPUB file.
        return (
            "<html xmlns='http://www.w3.org/1999/xhtml'><head>"
            f"<title>{html.escape(title)}</title></head><body>"
            + "".join(parts)
            + "</body></html>"
        )

    def _build_epub(self, book_data: Dict, chapters: List[Dict]) -> bytes:
        if not EPUB_AVAILABLE:
            raise ExportUnavailableError(
                "EPUB export unavailable: ebooklib is not installed"
            )

        book = epub.EpubBook()
        book.set_identifier(
            str(book_data.get("id") or book_data.get("_id") or "auto-author-book")
        )
        book.set_title(book_data.get("title", "Untitled"))
        book.set_language("en")

        author = book_data.get("author_name") or book_data.get("owner_name")
        if author:
            book.add_author(author)
        if book_data.get("description"):
            book.add_metadata("DC", "description", book_data["description"])
        if book_data.get("genre"):
            book.add_metadata("DC", "subject", book_data["genre"])

        # Title page
        meta_lines = []
        if book_data.get("subtitle"):
            meta_lines.append(f"<h2>{html.escape(book_data['subtitle'])}</h2>")
        if author:
            meta_lines.append(f"<p>by {html.escape(author)}</p>")
        for label, key in (("Genre", "genre"), ("Target Audience", "target_audience")):
            if book_data.get(key):
                meta_lines.append(f"<p>{label}: {html.escape(book_data[key])}</p>")
        if book_data.get("description"):
            meta_lines.append(f"<p>{html.escape(book_data['description'])}</p>")
        title_page = epub.EpubHtml(
            title="Title Page", file_name="title.xhtml", lang="en"
        )
        title_page.content = (
            "<html xmlns='http://www.w3.org/1999/xhtml'><head>"
            f"<title>{html.escape(book_data.get('title', 'Untitled'))}</title></head>"
            f"<body><h1>{html.escape(book_data.get('title', 'Untitled'))}</h1>"
            + "".join(meta_lines)
            + "</body></html>"
        )
        book.add_item(title_page)

        # Chapters
        epub_chapters = []
        for i, chapter in enumerate(chapters, 1):
            item = epub.EpubHtml(
                title=chapter.get("title", f"Chapter {i}"),
                file_name=f"chap_{i}.xhtml",
                lang="en",
            )
            item.content = self._chapter_to_xhtml(chapter, i)
            book.add_item(item)
            epub_chapters.append(item)

        # ponytail: flat nav — every chapter is a top-level nav point. EpubNcx +
        # EpubNav give working ereader navigation; nest by `level` if needed later.
        book.toc = [title_page, *epub_chapters]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = [title_page, "nav", *epub_chapters]

        output = io.BytesIO()
        epub.write_epub(output, book)
        output.seek(0)
        return output.getvalue()

    async def generate_markdown(
        self,
        book_data: Dict,
        chapters: List[Dict],
        multi_file: bool = False,
    ) -> bytes:
        """
        Generate Markdown from book data and chapters.

        ``multi_file=False`` returns a single ``.md`` file (UTF-8 bytes);
        ``multi_file=True`` returns a ZIP archive with one ``NN-slug.md`` per
        chapter. Runs in a worker thread so large books don't block the loop and
        an outer asyncio timeout can cancel — same pattern as PDF/DOCX/EPUB.
        """
        return await asyncio.to_thread(
            self._build_markdown, book_data, chapters, multi_file
        )

    @staticmethod
    def _slugify(text: str) -> str:
        """Lowercase, hyphenate, strip non-word chars for safe filenames."""
        slug = re.sub(r'[^\w\s-]', '', (text or '').lower()).strip()
        slug = re.sub(r'[\s_-]+', '-', slug).strip('-')
        return slug or 'chapter'

    def _content_to_markdown(self, content: str) -> str:
        """Convert chapter HTML to Markdown, preserving images and links.

        Uses a dedicated html2text instance with ``ignore_images=False`` so
        ``<img>`` becomes ``![alt](src)`` (issue #61 AC: images referenced with
        correct paths), unlike the shared text-extraction converter.
        """
        if not content:
            return "*(No content yet)*"
        if HTML2TEXT_AVAILABLE:
            h2t = html2text.HTML2Text()
            h2t.ignore_links = False
            h2t.ignore_images = False
            h2t.body_width = 0
            md = h2t.handle(content)
        else:
            # ponytail: no html2text — reuse the block-preserving tag strip.
            md = self._clean_html_content(content)
        return re.sub(r'\n{3,}', '\n\n', md).strip()

    def _chapter_markdown(self, chapter: Dict, index: int) -> str:
        """Render one chapter to Markdown (heading level tracks nesting depth)."""
        title = chapter.get('title', f'Chapter {index}')
        # level 1 chapter -> '##' (book title owns '#'); deeper nests indent.
        heading = '#' * min(6, chapter.get('level', 1) + 1)
        parts = [f"{heading} Chapter {index}: {title}"]
        if chapter.get('description'):
            parts.append(f"*{chapter['description']}*")
        parts.append(self._content_to_markdown(chapter.get('content', '')))
        return "\n\n".join(parts)

    def _build_markdown(
        self, book_data: Dict, chapters: List[Dict], multi_file: bool = False
    ) -> bytes:
        if not HTML2TEXT_AVAILABLE:
            raise ExportUnavailableError(
                "Markdown export unavailable: html2text is not installed"
            )

        if multi_file:
            output = io.BytesIO()
            import zipfile
            with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zf:
                for i, chapter in enumerate(chapters, 1):
                    name = f"{i:02d}-{self._slugify(chapter.get('title', ''))}.md"
                    zf.writestr(name, self._chapter_markdown(chapter, i) + "\n")
            output.seek(0)
            return output.getvalue()

        lines = [f"# {book_data.get('title', 'Untitled')}"]
        if book_data.get('subtitle'):
            # Emphasis, not a heading: the subtitle shouldn't compete with
            # chapter headings in a Markdown viewer's outline.
            lines.append(f"*{book_data['subtitle']}*")
        author = book_data.get('author_name') or book_data.get('owner_name')
        if author:
            lines.append(f"by {author}")
        metadata_parts = []
        if book_data.get('genre'):
            metadata_parts.append(f"Genre: {book_data['genre']}")
        if book_data.get('target_audience'):
            metadata_parts.append(f"Target Audience: {book_data['target_audience']}")
        if metadata_parts:
            lines.append(' • '.join(metadata_parts))
        if book_data.get('description'):
            lines.append(book_data['description'])

        for i, chapter in enumerate(chapters, 1):
            lines.append(self._chapter_markdown(chapter, i))

        return ("\n\n".join(lines).strip() + "\n").encode('utf-8')

    def _flatten_chapters(self, chapters: List[Dict], level: int = 1) -> List[Dict]:
        """Flatten nested chapter structure for export."""
        flattened = []

        for chapter in chapters:
            # Add the chapter itself
            chapter_copy = chapter.copy()
            chapter_copy['level'] = level
            flattened.append(chapter_copy)

            # Recursively add subchapters
            if chapter.get('subchapters'):
                flattened.extend(
                    self._flatten_chapters(chapter['subchapters'], level + 1)
                )

        return flattened

    def _select_chapters(
        self, book_data: Dict, include_empty_chapters: bool
    ) -> List[Dict]:
        """Flatten the TOC and drop empty chapters unless explicitly included."""
        toc = book_data.get('table_of_contents', {})
        flattened = self._flatten_chapters(toc.get('chapters', []))
        if include_empty_chapters:
            return flattened
        return [ch for ch in flattened if ch.get('content', '').strip()]

    def validate_book_for_export(
        self, book_data: Dict, include_empty_chapters: bool = False
    ) -> None:
        """
        Validate that a book has exportable content before starting generation.

        Raises ExportValidationError with a user-facing message when there is
        nothing to export, so the caller can return a clear 400 instead of
        producing a near-empty file or an opaque failure.
        """
        chapters = self._select_chapters(book_data, include_empty_chapters)
        if not chapters:
            toc = book_data.get('table_of_contents', {})
            has_chapters = bool(self._flatten_chapters(toc.get('chapters', [])))
            if has_chapters:
                raise ExportValidationError(
                    "This book has no chapter content yet. Add content to at "
                    "least one chapter, or enable 'Include empty chapters' to "
                    "export the outline."
                )
            raise ExportValidationError(
                "This book has no chapters to export. Generate a table of "
                "contents and add some content first."
            )

    def book_stats(self, book_data: Dict) -> Dict:
        """Compute export statistics (counts, word count, estimated pages)."""
        all_chapters = self._flatten_chapters(
            book_data.get('table_of_contents', {}).get('chapters', [])
        )
        with_content = [c for c in all_chapters if c.get('content', '').strip()]
        word_count = sum(
            len(self._clean_html_content(c.get('content', '')).split())
            for c in with_content
        )
        # ponytail: ~300 words/page is a fine rough estimate for a UI hint.
        estimated_pages = max(1, -(-word_count // 300)) if word_count else 0
        return {
            "total_chapters": len(all_chapters),
            "chapters_with_content": len(with_content),
            "total_word_count": word_count,
            "estimated_pages": estimated_pages,
        }

    async def export_book(
        self,
        book_data: Dict,
        format: str = "pdf",
        include_empty_chapters: bool = False,
        page_size: str = "letter",
        timeout_seconds: Optional[float] = None,
        template_id: Optional[str] = None,
        custom_options: Optional[Dict] = None,
        multi_file: bool = False,
    ) -> bytes:
        """
        Export a book in the specified format.

        Validates content first, then generates within a time budget. When a
        ``template_id`` is given, professional template styling is applied (with
        optional ``custom_options`` overrides). Raises ExportValidationError (no
        content), ExportTimeoutError (too slow), ValueError (bad format/options),
        or KeyError (unknown template).
        """
        self.validate_book_for_export(book_data, include_empty_chapters)

        flattened_chapters = self._select_chapters(book_data, include_empty_chapters)

        # Add author name from owner data if available
        if 'author_name' not in book_data and book_data.get('owner_name'):
            book_data['author_name'] = book_data['owner_name']

        template = resolve_template(template_id, custom_options)

        fmt = format.lower()
        if fmt == 'pdf':
            generator = self.generate_pdf(
                book_data, flattened_chapters, page_size=page_size, template=template
            )
        elif fmt == 'docx':
            generator = self.generate_docx(
                book_data, flattened_chapters, template=template
            )
        elif fmt == 'epub':
            # EPUB has its own reflowable layout — no page_size/template styling.
            generator = self.generate_epub(book_data, flattened_chapters)
        elif fmt in ('markdown', 'md'):
            # Markdown is plain text — no page_size/template styling.
            generator = self.generate_markdown(
                book_data, flattened_chapters, multi_file=multi_file
            )
        else:
            raise ValueError(f"Unsupported export format: {format}")

        if timeout_seconds is None:
            return await generator
        try:
            return await asyncio.wait_for(generator, timeout=timeout_seconds)
        except asyncio.TimeoutError as e:
            raise ExportTimeoutError(
                "Export took too long and was stopped. This usually happens "
                "with very large books — try exporting fewer chapters, or try "
                "again in a moment."
            ) from e


# Create singleton instance
export_service = ExportService()
