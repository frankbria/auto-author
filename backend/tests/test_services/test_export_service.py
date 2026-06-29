"""
Test Export Service functionality
"""
import asyncio
import pytest
from io import BytesIO
from unittest.mock import patch
from app.services.export_service import (
    export_service,
    ExportValidationError,
    ExportTimeoutError,
)
from reportlab.pdfgen import canvas
from docx import Document


class TestExportService:
    """Test the export service functionality."""

    @pytest.fixture
    def sample_book_data(self):
        """Sample book data for testing."""
        return {
            "title": "Test Book",
            "subtitle": "A Book for Testing",
            "description": "This is a test book used for testing the export functionality.",
            "genre": "Fiction",
            "target_audience": "General",
            "author_name": "Test Author",
            "table_of_contents": {
                "chapters": [
                    {
                        "id": "ch1",
                        "title": "Introduction",
                        "description": "The beginning of our story",
                        "content": "<h1>Welcome</h1><p>This is the <strong>first chapter</strong> of our book.</p><p>It contains some basic HTML formatting.</p>",
                        "order": 1,
                        "status": "completed",
                        "word_count": 100
                    },
                    {
                        "id": "ch2",
                        "title": "Chapter Two",
                        "description": "The plot thickens",
                        "content": "<h2>Chapter 2</h2><p>This chapter has a <em>different</em> heading level.</p><ul><li>Item 1</li><li>Item 2</li></ul>",
                        "order": 2,
                        "status": "completed",
                        "word_count": 150,
                        "subchapters": [
                            {
                                "id": "ch2.1",
                                "title": "Subsection 2.1",
                                "content": "<p>This is a subsection with its own content.</p>",
                                "order": 1,
                                "status": "draft",
                                "word_count": 50
                            }
                        ]
                    },
                    {
                        "id": "ch3",
                        "title": "Empty Chapter",
                        "description": "This chapter has no content",
                        "content": "",
                        "order": 3,
                        "status": "draft",
                        "word_count": 0
                    }
                ]
            }
        }

    def test_clean_html_content(self):
        """Test HTML to text conversion."""
        html = "<h1>Title</h1><p>This is a <strong>bold</strong> paragraph.</p>"
        result = export_service._clean_html_content(html)

        assert "# Title" in result
        assert "This is a **bold** paragraph." in result

    def test_extract_text_formatting(self):
        """Test text formatting extraction."""
        content = "<h1>Main Title</h1><p>Normal paragraph.</p><h2>Subtitle</h2>"
        formatted = export_service._extract_text_formatting(content)

        assert len(formatted) == 3
        assert formatted[0]['text'] == "Main Title"
        assert formatted[0]['style'] == 'heading1'
        assert formatted[1]['text'] == "Normal paragraph."
        assert formatted[1]['style'] == 'normal'
        assert formatted[2]['text'] == "Subtitle"
        assert formatted[2]['style'] == 'heading2'

    def test_flatten_chapters(self):
        """Test chapter flattening for nested structures."""
        chapters = [
            {
                "id": "1",
                "title": "Chapter 1",
                "subchapters": [
                    {"id": "1.1", "title": "Section 1.1"},
                    {"id": "1.2", "title": "Section 1.2"}
                ]
            },
            {"id": "2", "title": "Chapter 2"}
        ]

        flattened = export_service._flatten_chapters(chapters)

        assert len(flattened) == 4
        assert flattened[0]['level'] == 1
        assert flattened[1]['level'] == 2
        assert flattened[2]['level'] == 2
        assert flattened[3]['level'] == 1

    @pytest.mark.asyncio
    async def test_generate_pdf(self, sample_book_data):
        """Test PDF generation."""
        chapters = sample_book_data['table_of_contents']['chapters'][:2]  # Use first 2 chapters

        pdf_bytes = await export_service.generate_pdf(
            sample_book_data,
            chapters,
            page_size="letter"
        )

        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 1000  # Should have substantial content
        assert pdf_bytes.startswith(b'%PDF')  # PDF magic number

    @pytest.mark.asyncio
    async def test_generate_docx(self, sample_book_data):
        """Test DOCX generation."""
        chapters = sample_book_data['table_of_contents']['chapters'][:2]  # Use first 2 chapters

        docx_bytes = await export_service.generate_docx(
            sample_book_data,
            chapters
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000  # Should have substantial content
        # DOCX files start with PK (ZIP format)
        assert docx_bytes.startswith(b'PK')

    @pytest.mark.asyncio
    async def test_export_book_pdf(self, sample_book_data):
        """Test complete book export to PDF."""
        pdf_bytes = await export_service.export_book(
            sample_book_data,
            format="pdf",
            include_empty_chapters=False
        )

        assert isinstance(pdf_bytes, bytes)
        assert len(pdf_bytes) > 1000
        assert pdf_bytes.startswith(b'%PDF')

    @pytest.mark.asyncio
    async def test_export_book_docx(self, sample_book_data):
        """Test complete book export to DOCX."""
        docx_bytes = await export_service.export_book(
            sample_book_data,
            format="docx",
            include_empty_chapters=True
        )

        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000
        assert docx_bytes.startswith(b'PK')

    @pytest.mark.asyncio
    async def test_export_exclude_empty_chapters(self, sample_book_data):
        """Test that empty chapters are excluded when requested."""
        # Export without empty chapters
        pdf_bytes = await export_service.export_book(
            sample_book_data,
            format="pdf",
            include_empty_chapters=False
        )

        # The PDF should not contain "Empty Chapter" text
        # This is a basic check - in reality you'd parse the PDF
        assert b"Empty Chapter" not in pdf_bytes

    @pytest.mark.asyncio
    async def test_export_invalid_format(self, sample_book_data):
        """Test that invalid format raises error."""
        with pytest.raises(ValueError, match="Unsupported export format"):
            await export_service.export_book(
                sample_book_data,
                format="invalid"
            )

    @pytest.mark.asyncio
    async def test_export_with_no_content(self):
        """A book with no chapters is rejected with a clear validation error."""
        minimal_book = {
            "title": "Empty Book",
            "table_of_contents": {
                "chapters": []
            }
        }

        with pytest.raises(ExportValidationError) as exc:
            await export_service.export_book(minimal_book, format="pdf")
        assert "no chapters" in str(exc.value).lower()

    @pytest.mark.asyncio
    async def test_export_empty_chapters_rejected_unless_included(self):
        """Chapters that exist but have no content are rejected by default,
        but exportable when include_empty_chapters is set."""
        book = {
            "title": "Outline Only",
            "table_of_contents": {
                "chapters": [
                    {"id": "1", "title": "Ch One", "content": "", "order": 1}
                ]
            },
        }

        with pytest.raises(ExportValidationError) as exc:
            await export_service.export_book(book, format="pdf")
        assert "no chapter content" in str(exc.value).lower()

        pdf_bytes = await export_service.export_book(
            book, format="pdf", include_empty_chapters=True
        )
        assert pdf_bytes.startswith(b"%PDF")

    @pytest.mark.asyncio
    async def test_book_stats_counts_words_and_pages(self):
        """book_stats reports chapter counts, word count, and estimated pages."""
        book = {
            "title": "Stats Book",
            "table_of_contents": {
                "chapters": [
                    {"id": "1", "title": "A",
                     "content": "<p>one two three four five</p>", "order": 1},
                    {"id": "2", "title": "B", "content": "", "order": 2},
                ]
            },
        }
        stats = export_service.book_stats(book)
        assert stats["total_chapters"] == 2
        assert stats["chapters_with_content"] == 1
        assert stats["total_word_count"] == 5
        assert stats["estimated_pages"] >= 1

    @pytest.mark.asyncio
    async def test_export_times_out_for_slow_generation(self):
        """A generation that exceeds the timeout raises ExportTimeoutError."""
        book = {
            "title": "Slow Book",
            "table_of_contents": {
                "chapters": [
                    {"id": "1", "title": "A", "content": "<p>hi</p>", "order": 1}
                ]
            },
        }

        async def slow_pdf(*args, **kwargs):
            await asyncio.sleep(0.2)
            return b"%PDF-never"

        with patch.object(export_service, "generate_pdf", side_effect=slow_pdf):
            with pytest.raises(ExportTimeoutError):
                await export_service.export_book(
                    book, format="pdf", timeout_seconds=0.01
                )

    @pytest.mark.asyncio
    async def test_export_large_book_succeeds(self):
        """A large book (100+ chapters) exports successfully within budget."""
        chapters = [
            {
                "id": str(i),
                "title": f"Chapter {i}",
                "content": f"<p>Content for chapter {i} with several words.</p>",
                "order": i,
            }
            for i in range(1, 121)
        ]
        book = {"title": "Big Book", "table_of_contents": {"chapters": chapters}}

        pdf_bytes = await export_service.export_book(
            book, format="pdf", timeout_seconds=60
        )
        assert pdf_bytes.startswith(b"%PDF")
        assert len(pdf_bytes) > 1000

    @pytest.mark.asyncio
    async def test_export_special_characters(self):
        """Test export with special characters in content."""
        book_data = {
            "title": "Special Characters & Symbols",
            "author_name": "Author <Name>",
            "table_of_contents": {
                "chapters": [{
                    "id": "1",
                    "title": "Symbols & More",
                    "content": "<p>Special chars: &amp; &lt; &gt; &quot;</p>",
                    "order": 1
                }]
            }
        }

        # Should not raise any errors
        pdf_bytes = await export_service.export_book(book_data, format="pdf")
        assert isinstance(pdf_bytes, bytes)

        docx_bytes = await export_service.export_book(book_data, format="docx")
        assert isinstance(docx_bytes, bytes)

class TestExportAvailabilityGuards:
    """Issue #45: export degrades gracefully when an optional library is missing."""

    @pytest.mark.asyncio
    async def test_generate_pdf_raises_when_reportlab_unavailable(self, monkeypatch):
        import app.services.export_service as es
        monkeypatch.setattr(es, "PDF_AVAILABLE", False)
        with pytest.raises(es.ExportUnavailableError, match="reportlab"):
            await es.export_service.generate_pdf({"title": "x"}, [])

    @pytest.mark.asyncio
    async def test_generate_docx_raises_when_docx_unavailable(self, monkeypatch):
        import app.services.export_service as es
        monkeypatch.setattr(es, "DOCX_AVAILABLE", False)
        with pytest.raises(es.ExportUnavailableError, match="python-docx"):
            await es.export_service.generate_docx({"title": "x"}, [])

    def test_clean_html_falls_back_without_html2text(self):
        import app.services.export_service as es
        svc = es.ExportService()
        svc.h2t = None  # simulate html2text not installed
        out = svc._clean_html_content("<p>Hello <strong>world</strong></p>")
        assert "Hello" in out and "world" in out
        assert "<" not in out


class TestTemplatedExport:
    """Export with professional templates (issue #59)."""

    @pytest.fixture
    def book_data(self):
        return {
            "title": "Templated Book",
            "subtitle": "A Test",
            "author_name": "Jane Doe",
            "description": "Desc.",
            "table_of_contents": {
                "chapters": [
                    {
                        "id": "c1",
                        "title": "One",
                        "content": "<p>Some real content for the first chapter.</p>",
                        "order": 1,
                        "status": "completed",
                        "word_count": 8,
                    }
                ]
            },
        }

    @pytest.mark.parametrize(
        "template_id", ["classic_fiction", "modern_nonfiction", "academic"]
    )
    @pytest.mark.asyncio
    async def test_pdf_export_with_each_template(self, book_data, template_id):
        out = await export_service.export_book(
            book_data, format="pdf", template_id=template_id
        )
        assert out[:4] == b"%PDF"
        assert len(out) > 500

    @pytest.mark.asyncio
    async def test_pdf_templates_differ_from_default(self, book_data):
        default = await export_service.export_book(book_data, format="pdf")
        templated = await export_service.export_book(
            book_data, format="pdf", template_id="academic"
        )
        assert default != templated  # different page size/margins/font

    @pytest.mark.asyncio
    async def test_docx_template_applies_section_and_font(self, book_data):
        out = await export_service.export_book(
            book_data, format="docx", template_id="academic"
        )
        doc = Document(BytesIO(out))
        section = doc.sections[0]
        # Academic = A4, 1.25" inside/outside, 1.0" top/bottom.
        assert round(section.left_margin.inches, 2) == 1.25
        assert round(section.right_margin.inches, 2) == 1.25
        assert round(section.top_margin.inches, 2) == 1.0
        normal = doc.styles["Normal"]
        assert normal.font.name == "Times New Roman"
        assert normal.font.size.pt == 12
        # Running header carries book/author text.
        assert "Templated Book" in section.header.paragraphs[0].text

    @pytest.mark.asyncio
    async def test_custom_options_override_font_size(self, book_data):
        out = await export_service.export_book(
            book_data,
            format="docx",
            template_id="classic_fiction",
            custom_options={"font_size": 14},
        )
        doc = Document(BytesIO(out))
        assert doc.styles["Normal"].font.size.pt == 14

    @pytest.mark.asyncio
    async def test_unknown_template_raises(self, book_data):
        with pytest.raises(KeyError):
            await export_service.export_book(
                book_data, format="pdf", template_id="not_a_template"
            )

    @pytest.mark.asyncio
    async def test_no_template_id_is_backward_compatible(self, book_data):
        out = await export_service.export_book(book_data, format="docx")
        doc = Document(BytesIO(out))
        # Default python-docx letter section, untouched header.
        assert doc.sections[0].header.paragraphs[0].text == ""
